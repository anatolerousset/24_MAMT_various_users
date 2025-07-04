import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Optional, Union, List, Set

import textwrap

from datetime import datetime


from docling_core.types.doc import (
    DocItemLabel,
    DoclingDocument,
    DocumentOrigin,
    GroupLabel,
    ImageRef,
    NodeItem,
    TableCell,
    TableData,
)


from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tc
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree
from lxml.etree import XPath
from PIL import Image, UnidentifiedImageError
from typing_extensions import override

from docling.backend.abstract_backend import DeclarativeDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument


from docling_core.types.doc.labels import CodeLanguageLabel, DocItemLabel, GroupLabel
from docling_core.types.doc import SectionHeaderItem, TextItem

_log = logging.getLogger(__name__)

from docling_core.types.doc import DocumentOrigin

class DocumentOriginMod(DocumentOrigin):
    """FileSource with additional metadata."""
    creation_date: Optional[datetime] = None  
    last_modified_date: Optional[datetime] = None  # Last modified date of the document
    author: Optional[str] = None  # Author of the document


class MsWordDocumentBackendModMetadataplus(DeclarativeDocumentBackend):
    @override
    def __init__(
        self, in_doc: "InputDocument", path_or_stream: Union[BytesIO, Path]
    ) -> None:
        super().__init__(in_doc, path_or_stream)
        self.XML_KEY = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )
        self.xml_namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
            "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
            "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
            "dgm1": "http://schemas.microsoft.com/office/drawing/2007/8/2/diagram",
            "dsp": "http://schemas.microsoft.com/office/drawing/2008/diagram",
            "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
            "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
            "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",

            "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
            "dc": "http://purl.org/dc/elements/1.1/",
            "dcterms": "http://purl.org/dc/terms/",
            "xsi": "http://www.w3.org/2001/XMLSchema-instance",
        }
        
        # Keep track of the current list structure state
        self._list_last_level_seen = None
        self._current_list_depth = 0
        # Word file:
        self.path_or_stream: Union[BytesIO, Path] = path_or_stream
        self.valid: bool = False
        # Initialise the parents for the hierarchy
        self.max_levels: int = 10
        self.level_at_new_list: Optional[int] = None
        self.parents: dict[int, Optional[NodeItem]] = {}
        for i in range(-1, self.max_levels):
            self.parents[i] = None

        self.level = 0
        self.listIter = 0
        self.current_main_bullet = None  # Track the current main bullet item

        self.history: dict[str, Any] = {
            "names": [None],
            "levels": [None],
            "numids": [None],
            "indents": [None],
        }
        
        # Store images found in tables to process after the table
        self.pending_table_images: List[tuple[Any, Optional[NodeItem]]] = []

        # Define multilingual style mappings
        self.title_styles: Set[str] = {
            "Title", "Titredulivre", "TitreduLivre", "BookTitle", 
            "MainTitle", "TitrePrincipal"
        }
        
        self.heading_patterns = [
            re.compile(r"^heading\s*(\d+)$", re.IGNORECASE),
            re.compile(r"^titre\s*(\d+)$", re.IGNORECASE),
            re.compile(r"^heading(\d+)$", re.IGNORECASE),
            re.compile(r"^titre(\d+)$", re.IGNORECASE),
            re.compile(r"^h(\d+)$", re.IGNORECASE),
            re.compile(r"^titreniveau(\d+)$", re.IGNORECASE)
        ]
        
        self.bullet_styles: Set[str] = {
            "Listepuces", "Listepuces2", "BulletList", "BulletedList", 
            "ListBullet", "ListBullet2"
        }
        
        self.paragraph_styles: Set[str] = {
            "Paragraph", "Paragraphe", "Normal", "Subtitle", "SousTitre", 
            "Author", "Auteur", "DefaultText", "TexteParDefaut", 
            "ListParagraph", "ParagrapheListe", "Quote", "Citation"
        }

        self.docx_obj = None
        try:
            if isinstance(self.path_or_stream, BytesIO):
                self.docx_obj = Document(self.path_or_stream)
            elif isinstance(self.path_or_stream, Path):
                self.docx_obj = Document(str(self.path_or_stream))

            self.valid = True
        except Exception as e:
            raise RuntimeError(
                f"MsPowerpointDocumentBackend could not load document with hash {self.document_hash}"
            ) from e

    @override
    def is_valid(self) -> bool:
        return self.valid

    @classmethod
    @override
    def supports_pagination(cls) -> bool:
        return False

    @override
    def unload(self):
        if isinstance(self.path_or_stream, BytesIO):
            self.path_or_stream.close()

        self.path_or_stream = None

    @classmethod
    @override
    def supported_formats(cls) -> set[InputFormat]:
        return {InputFormat.DOCX}

    @override
    def convert(self) -> DoclingDocument:
        """Parses the DOCX into a structured document model.

        Returns:
            The parsed document.
        """
        # Extract creation and modification dates, and author
        created_date, modified_date, author = self.extract_document_metadata()
        #print(created_date, modified_date, author)

        origin = DocumentOriginMod(
            filename=self.file.name or "file",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash=self.document_hash,
            creation_date=created_date,    
            last_modified_date=modified_date,
            author=author
        )

        doc = DoclingDocument(name=self.file.stem or "file", origin=origin)
        if self.is_valid():
            assert self.docx_obj is not None
            doc = self.walk_linear(self.docx_obj.element.body, self.docx_obj, doc)
            return doc
        else:
            raise RuntimeError(
                f"Cannot convert doc with {self.document_hash} because the backend failed to init."
            )
    def update_history(
        self,
        name: str,
        level: Optional[int],
        numid: Optional[int],
        ilevel: Optional[int],
    ):
        self.history["names"].append(name)
        self.history["levels"].append(level)
        self.history["numids"].append(numid)
        self.history["indents"].append(ilevel)

    def prev_name(self) -> Optional[str]:
        return self.history["names"][-1]

    def prev_level(self) -> Optional[int]:
        return self.history["levels"][-1]

    def prev_numid(self) -> Optional[int]:
        return self.history["numids"][-1]

    def prev_indent(self) -> Optional[int]:
        return self.history["indents"][-1]

    def prev_style_id(self) -> Optional[str]:
        return self.history["names"][-1]

    def get_level(self) -> int:
        """Return the first None index."""
        for k, v in self.parents.items():
            if k >= 0 and v == None:
                return k
        return 0

    def walk_linear(
        self,
        body: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
        inside_table = False,
    ) -> DoclingDocument:
        for element in body:
            tag_name = etree.QName(element).localname
            
            # Check for Inline Images (blip elements)
            xpath_expr = XPath(".//a:blip", namespaces=self.xml_namespaces)
            drawing_blip = xpath_expr(element)

            # Check for text boxes (wps:txbx element)
            xpath_textbox = XPath(".//wps:txbx/w:txbxContent", namespaces=self.xml_namespaces)
            textbox_content = xpath_textbox(element)

            # Check for diagrams
            xpath_diagram = XPath(".//a:graphicData[@uri='http://schemas.openxmlformats.org/drawingml/2006/diagram']", 
                            namespaces=self.xml_namespaces)
            diagram_data = xpath_diagram(element)

            if not inside_table and diagram_data:
                self.handle_diagram(element, docx_obj, doc)

            # Check for Tables
            if element.tag.endswith("tbl"):
                try:
                    # Save current parent level for post-table processing
                    current_parent_level = self.get_level() - 1
                    current_parent = self.parents[current_parent_level]
                    
                    # Process the table content
                    self.handle_tables(element, docx_obj, doc)
                    
                    # Clear the pending images list after processing
                    self.pending_table_images = []
                except Exception as e:
                    _log.debug(f"Could not parse a table, broken docx table: {str(e)}")

            if not inside_table and drawing_blip:
                self.handle_pictures(docx_obj, drawing_blip, doc)

            # Handle text boxes
            if textbox_content:
                for txbx in textbox_content:
                    self.handle_textbox(txbx, docx_obj, doc)
            
            # Check for the sdt containers, like table of contents
            elif tag_name in ["sdt"]:
                # First check if this is a Table of Contents we want to skip
                if self.is_table_of_contents(element):
                    _log.debug("Skipping Table of Contents SDT element")
                    continue  # Skip this element
                    
                # Process other SDT content that isn't a TOC
                sdt_content = element.find(".//w:sdtContent", namespaces=self.xml_namespaces)
                if sdt_content is not None:
                    # Iterate paragraphs, runs, or text inside <w:sdtContent>.
                    paragraphs = sdt_content.findall(".//w:p", namespaces=self.xml_namespaces)
                    for p in paragraphs:
                        self.handle_text_elements(p, docx_obj, doc)
                        
            # Check for Text
            elif tag_name in ["p"]:
                # "tcPr", "sectPr"
                self.handle_text_elements(element, docx_obj, doc)
            else:
                _log.debug(f"Ignoring element in DOCX with tag: {tag_name}")
        return doc

    # Add a new method to identify if an SDT element is a table of contents
    def is_table_of_contents(self, element: BaseOxmlElement) -> bool:
        """
        Determine if an SDT element is a Table of Contents.
        
        Args:
            element: The SDT element to check
            
        Returns:
            bool: True if the element is a Table of Contents, False otherwise
        """
        # Check for TOC by looking at the sdtPr (SDT Properties)
        sdt_props = element.find(".//w:sdtPr", namespaces=self.xml_namespaces)
        if sdt_props is not None:
            # Method 1: Check for TOC docPartObj
            doc_part = sdt_props.find(".//w:docPartObj/w:docPartGallery", namespaces=self.xml_namespaces)
            if doc_part is not None:
                # Check the val attribute
                val = doc_part.get(self.XML_KEY)
                if val and "Table of Contents" in val:
                    return True
            
            # Method 2: Check for TOC alias
            alias = sdt_props.find(".//w:alias", namespaces=self.xml_namespaces)
            if alias is not None:
                val = alias.get(self.XML_KEY)
                if val and ("TOC" in val or "Table of Contents" in val or "TableOfContents" in val):
                    return True
                    
            # Method 3: Check for TOC tag
            tag = sdt_props.find(".//w:tag", namespaces=self.xml_namespaces)
            if tag is not None:
                val = tag.get(self.XML_KEY)
                if val and ("TOC" in val or "Table of Contents" in val):
                    return True
                    
            # Method 4: Check if it contains a specific "TOC" style
            styles = element.findall(".//w:pStyle", namespaces=self.xml_namespaces)
            for style in styles:
                val = style.get(self.XML_KEY, "")
                if val and ("TOC" in val or "TableOfContents" in val):
                    return True
            
            # Method 5: Check for the TOC field (w:fldChar w:fldCharType="begin" followed by w:instrText containing "TOC")
            field_char = element.find(".//w:fldChar[@w:fldCharType='begin']", namespaces=self.xml_namespaces)
            if field_char is not None:
                # Look for instruction text nodes that might contain TOC instructions
                instr_elements = element.findall(".//w:instrText", namespaces=self.xml_namespaces)
                for instr in instr_elements:
                    if instr.text and "TOC" in instr.text:
                        return True
        
        return False

    def str_to_int(self, s: Optional[str], default: Optional[int] = 0) -> Optional[int]:
        if s is None:
            return None
        try:
            return int(s)
        except ValueError:
            return default

    def extract_heading_level(self, style_id: str) -> Optional[int]:
        """Extract heading level from style ID using multiple patterns for different languages."""
        for pattern in self.heading_patterns:
            match = pattern.search(style_id)
            if match:
                try:
                    return int(match.group(1))
                except (ValueError, IndexError):
                    continue
        return None
        
    def get_outline_level(self, paragraph: Paragraph) -> Optional[int]:
        """Extract outline level from paragraph properties."""
        # Use the defined namespace dictionary for the find operation
        outline_lvl = paragraph._element.find(
            ".//w:outlineLvl", namespaces=self.xml_namespaces
        )

        if outline_lvl is not None:
            # Replace print statements with proper logging if needed
            # _log.debug(f"Found outline level element: {outline_lvl}")
            #print(paragraph.text)
            
            # Use the defined XML_KEY for consistency
            level_val = outline_lvl.get(self.XML_KEY)
            # _log.debug(f"Outline level value: {level_val}")
            try:
                return int(level_val)
            except (ValueError, TypeError):
                return None
        
        return None

    def get_numId_and_ilvl(
        self, paragraph: Paragraph
    ) -> tuple[Optional[int], Optional[int]]:
        # Access the XML element of the paragraph
        numPr = paragraph._element.find(
            ".//w:numPr", namespaces=paragraph._element.nsmap
        )

        if numPr is not None:
            # Get the numId element and extract the value
            numId_elem = numPr.find("w:numId", namespaces=paragraph._element.nsmap)
            ilvl_elem = numPr.find("w:ilvl", namespaces=paragraph._element.nsmap)
            numId = numId_elem.get(self.XML_KEY) if numId_elem is not None else None
            ilvl = ilvl_elem.get(self.XML_KEY) if ilvl_elem is not None else None

            return self.str_to_int(numId, None), self.str_to_int(ilvl, None)

        return None, None  # If the paragraph is not part of a list

    def get_label_and_level(self, paragraph: Paragraph) -> tuple[str, Optional[int]]:
        if paragraph.style is None:
            return "Normal", None
        
        label = paragraph.style.style_id
        if label is None:
            return "Normal", None
            
        # Check for heading patterns in various languages
        heading_level = self.extract_heading_level(label)
        if heading_level is not None:
            # Standardize to "Heading" for internal processing
            return "Heading", heading_level
            
        # Check for other style patterns like "Niveau" (French for "Level")
        if "Niveau" in label:
            parts = label.split("Niveau")
            if len(parts) == 2:
                return parts[0], self.str_to_int(parts[1], None)
                
        # Handle colon-separated level indicators (e.g., "SomeStyle:3")
        if ":" in label:
            parts = label.split(":")
            if len(parts) == 2:
                return parts[0], self.str_to_int(parts[1], None)

        # Fallback to the original style ID with no level
        return label, None

    def handle_text_elements(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        paragraph = Paragraph(element, docx_obj)

        if paragraph.text is None:
            return
        text = paragraph.text.strip()
        # Remove all newline characters (both \n and \r\n)
        text = text.replace('\r\n', ' ').replace('\n', ' ')
        # Replace multiple spaces with a single space
        text = ' '.join(text.split())

        # Check for images in the paragraph
        xpath_expr = XPath(".//a:blip", namespaces=self.xml_namespaces)
        drawing_blip = xpath_expr(element)

        # Check for outline level
        outline_level = self.get_outline_level(paragraph)

        # Identify whether list is a numbered list or not
        is_numbered = False
        p_style_id, p_level = self.get_label_and_level(paragraph)
        numid, ilevel = self.get_numId_and_ilvl(paragraph)

        if numid == 0:
            numid = None

        # Normalize style ID for case-insensitive comparison
        normalized_style = p_style_id.lower() if p_style_id else ""
        
        # Handle bullet lists in different languages
        is_bullet_style = any(bullet_style.lower() == normalized_style for bullet_style in self.bullet_styles)
        is_main_bullet = False
        is_sub_bullet = False

        # Determine if this is a main bullet or sub-bullet style
        if is_bullet_style:
            # Check if this is a main bullet style (first level)
            is_main_bullet = any(
                bullet_style.lower() == normalized_style and not "2" in bullet_style.lower()
                for bullet_style in self.bullet_styles
            )
            
            # Check if this is a sub-bullet style (second level)
            is_sub_bullet = any(
                bullet_style.lower() == normalized_style and "2" in bullet_style.lower()
                for bullet_style in self.bullet_styles
            )
            
            level = self.get_level()
            
            # Handle main bullet points
            if is_main_bullet:
                # Track the current main bullet item for proper sub-bullet attachment
                self.current_main_bullet = text
                
                # If this is the first bullet point or we're transitioning from a non-bullet style
                prev_style = self.prev_style_id()
                is_prev_bullet = prev_style and any(bullet.lower() in prev_style.lower() for bullet in self.bullet_styles)
                
                if not is_prev_bullet:
                    # Reset list structure completely
                    self.level_at_new_list = level
                    # Clear any existing lists
                    for key in range(len(self.parents)):
                        if key >= level:
                            self.parents[key] = None
                    
                    # Create new main list
                    self.parents[level] = doc.add_group(
                        label=GroupLabel.LIST, name="list", parent=self.parents[level - 1]
                    )
                
                # If coming from a sub-bullet point, make sure we're still at the main list
                elif is_prev_bullet and "2" in prev_style:
                    # Clear any existing sublist
                    if self.level_at_new_list is not None:
                        self.parents[self.level_at_new_list + 1] = None
                
                # Add the bullet point to the main list
                if text != "":
                    doc.add_list_item(
                        marker="",  # No explicit marker, the styling is handled by the document
                        enumerated=False,
                        parent=self.parents[self.level_at_new_list] if self.level_at_new_list is not None else self.parents[level],
                        text=text,
                    )
            
            # Handle sub-bullet points
            elif is_sub_bullet:
                # If there's no main list established yet, something is wrong
                # We'll create one just to be safe
                if self.level_at_new_list is None:
                    self.level_at_new_list = level
                    self.parents[level] = doc.add_group(
                        label=GroupLabel.LIST, name="list", parent=self.parents[level - 1]
                    )
                
                # Create a sublist if we're transitioning from a main bullet point
                # or if there's no existing sublist
                if self.parents[self.level_at_new_list + 1] is None:
                    self.parents[self.level_at_new_list + 1] = doc.add_group(
                        label=GroupLabel.LIST, 
                        name="sublist", 
                        parent=self.parents[self.level_at_new_list]
                    )
                
                # Add the sub-bullet point to the sublist
                if text != "":
                    doc.add_list_item(
                        marker="",  # No explicit marker
                        enumerated=False,
                        parent=self.parents[self.level_at_new_list + 1],
                        text=text,
                    )
            
            # Process any images in the paragraph after the list item
            if drawing_blip:
                parent_level = self.level_at_new_list if self.level_at_new_list is not None else level
                parent = self.parents[parent_level]
                self.handle_pictures(docx_obj, drawing_blip, doc, parent=parent)
                
            self.update_history(p_style_id, p_level, numid, ilevel)
            return

        # If we're coming from a list context but now have a non-list paragraph,
        # close the list structure
        prev_style = self.prev_style_id()
        prev_is_bullet = prev_style and any(bullet.lower() in prev_style.lower() for bullet in self.bullet_styles)

        if prev_is_bullet and not is_bullet_style:
            if self.level_at_new_list is not None:
                # Reset list-related structure
                for key in range(len(self.parents)):
                    if key >= self.level_at_new_list:
                        self.parents[key] = None
                self.level_at_new_list = None
                # Reset current main bullet
                if hasattr(self, 'current_main_bullet'):
                    self.current_main_bullet = None
        
        # Handle the original list structure with numid and ilevel
        elif (
            numid is not None
            and ilevel is not None
            and not any(p_style_id.lower() == title.lower() for title in self.title_styles)
            and p_style_id != "Heading"
            and outline_level is None
        ):
            self.add_listitem(
                doc,
                numid,
                ilevel,
                text,
                is_numbered,
            )
            self.update_history(p_style_id, p_level, numid, ilevel)
            return
            
        elif (
            numid is None
            and self.prev_numid() is not None
            and not any(p_style_id.lower() == title.lower() for title in self.title_styles)
            and p_style_id != "Heading"
            and outline_level is None
        ):  # Close list
            if self.level_at_new_list:
                for key in range(len(self.parents)):
                    if key >= self.level_at_new_list:
                        self.parents[key] = None
                self.level = self.level_at_new_list - 1
                self.level_at_new_list = None
                # Reset current main bullet
                self.current_main_bullet = None
            else:
                for key in range(len(self.parents)):
                    self.parents[key] = None
                self.level = 0
                # Reset current main bullet
                self.current_main_bullet = None

        # Handle title styles in any language
        if any(p_style_id.lower() == title.lower() for title in self.title_styles):
            for key in range(len(self.parents)):
                self.parents[key] = None
            self.parents[0] = doc.add_text(
                parent=None, label=DocItemLabel.TITLE, text=text
            )
        
        # Handle heading styles (now standardized to "Heading" with level)
        elif p_style_id == "Heading" and p_level is not None:
            self.add_header(doc, p_level, text)
            
        # Handle paragraphs with outline level (treat as headings)
        elif outline_level is not None:
            self.add_header(doc, outline_level + 1, text)  # Add 1 to match heading levels (0-based to 1-based)
            
        # Handle paragraph styles in any language
        elif any(p_style_id.lower() == para.lower() for para in self.paragraph_styles):
            level = self.get_level()
            doc.add_text(
                label=DocItemLabel.PARAGRAPH, parent=self.parents[level - 1], text=text
            )
        else:
            # Text style names can vary by language and user customization
            # So treat all other labels as paragraphs
            level = self.get_level()
            doc.add_text(
                label=DocItemLabel.PARAGRAPH, parent=self.parents[level - 1], text=text
            )
                
        self.update_history(p_style_id, p_level, numid, ilevel)
        return
    

    def add_header(
        self, doc: DoclingDocument, curr_level: Optional[int], text: str
    ) -> None:
        # Skip adding the heading if the text is empty
        if not text:
            return
            
        level = self.get_level()
        if isinstance(curr_level, int):
            if curr_level > level:
                # add invisible group
                for i in range(level, curr_level):
                    self.parents[i] = doc.add_group(
                        parent=self.parents[i - 1],
                        label=GroupLabel.SECTION,
                        name=f"header-{i}",
                    )
            elif curr_level < level:
                # remove the tail
                for key in range(len(self.parents)):
                    if key >= curr_level:
                        self.parents[key] = None

            self.parents[curr_level] = doc.add_heading(
                parent=self.parents[curr_level - 1],
                text=text,
                level=curr_level,
            )
        else:
            self.parents[self.level] = doc.add_heading(
                parent=self.parents[self.level - 1],
                text=text,
                level=1,
            )
        return

    def add_listitem(
        self,
        doc: DoclingDocument,
        numid: int,
        ilevel: int,
        text: str,
        is_numbered: bool = False,
    ) -> None:
        enum_marker = ""

        level = self.get_level()
        prev_indent = self.prev_indent()
        if self.prev_numid() is None:  # Open new list
            self.level_at_new_list = level

            self.parents[level] = doc.add_group(
                label=GroupLabel.LIST, name="list", parent=self.parents[level - 1]
            )

            # Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            doc.add_list_item(
                marker=enum_marker,
                enumerated=is_numbered,
                parent=self.parents[level],
                text=text,
            )

        elif (
            self.prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and prev_indent < ilevel
        ):  # Open indented list
            for i in range(
                self.level_at_new_list + prev_indent + 1,
                self.level_at_new_list + ilevel + 1,
            ):
                # Determine if this is an unordered list or an ordered list.
                # Set GroupLabel.ORDERED_LIST when it fits.
                self.listIter = 0
                if is_numbered:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.ORDERED_LIST,
                        name="list",
                        parent=self.parents[i - 1],
                    )
                else:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.LIST, name="list", parent=self.parents[i - 1]
                    )

            # Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            doc.add_list_item(
                marker=enum_marker,
                enumerated=is_numbered,
                parent=self.parents[self.level_at_new_list + ilevel],
                text=text,
            )

        elif (
            self.prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and ilevel < prev_indent
        ):  # Close list
            for k, v in self.parents.items():
                if k > self.level_at_new_list + ilevel:
                    self.parents[k] = None

            # Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            doc.add_list_item(
                marker=enum_marker,
                enumerated=is_numbered,
                parent=self.parents[self.level_at_new_list + ilevel],
                text=text,
            )
            self.listIter = 0

        elif self.prev_numid() == numid or prev_indent == ilevel:
            # Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            doc.add_list_item(
                marker=enum_marker,
                enumerated=is_numbered,
                parent=self.parents[level - 1],
                text=text,
            )
        return

    def handle_tables(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        table: Table = Table(element, docx_obj)
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        _log.debug(f"Table grid with {num_rows} rows and {num_cols} columns")
        
        # Clear the pending table images list before processing the table
        self.pending_table_images = []
        
        # Find all images in the table and store them for later processing
        xpath_expr = XPath(".//a:blip", namespaces=self.xml_namespaces)
        table_images = xpath_expr(element)
        if table_images:
            level = self.get_level()
            parent = self.parents[level - 1]
            for blip in table_images:
                self.pending_table_images.append((blip, parent))

        # Check for header styles (now checking both "Titre3/4" and English equivalents)
        #first_cell_style = table.rows[0].cells[0].paragraphs[0].style.style_id if table.rows and table.rows[0].cells else None
        # Check for header styles in first and second cells
        first_cell_style = None
        if table.rows and table.rows[0].cells:
            # Check first cell
            if len(table.rows[0].cells) > 0 and table.rows[0].cells[0].paragraphs:
                first_cell_style = table.rows[0].cells[0].paragraphs[0].style.style_id
            
            # If no header style found in first cell, check second cell
            if (first_cell_style is None or not any(pattern.search(first_cell_style) for pattern in self.heading_patterns)) and len(table.rows[0].cells) > 1:
                if table.rows[0].cells[1].paragraphs:
                    first_cell_style = table.rows[0].cells[1].paragraphs[0].style.style_id
                    
        # Extract heading level from the style if it's a heading
        heading_level = None
        if first_cell_style:
            for pattern in self.heading_patterns:
                match = pattern.search(first_cell_style)
                if match:
                    try:
                        heading_level = int(match.group(1))
                        break
                    except (ValueError, IndexError):
                        continue
        
        # Process table as sequences of paragraphs if it has heading cells or is a single-row/column table
        if (heading_level and heading_level >= 3) or num_rows <=  2 and num_cols <= 4:
            for row_idx, row in enumerate(table.rows):
                _log.debug(f"Row index {row_idx} with {len(row.cells)} populated cells")
                col_idx = 0
                while col_idx < num_cols:
                    cell_element = table.rows[row_idx].cells[col_idx]
                    col_idx += 1
                    self.walk_linear(cell_element._element, docx_obj, doc,inside_table=True)
            return
            
        else:
            data = TableData(num_rows=num_rows, num_cols=num_cols)
            cell_set: set[CT_Tc] = set()
            for row_idx, row in enumerate(table.rows):
                _log.debug(f"Row index {row_idx} with {len(row.cells)} populated cells")
                col_idx = 0
                while col_idx < num_cols:
                    cell: _Cell = row.cells[col_idx]
                    _log.debug(
                        f" col {col_idx} grid_span {cell.grid_span} grid_cols_before {row.grid_cols_before}"
                    )
                    if cell is None or cell._tc in cell_set:
                        _log.debug(f"  skipped since repeated content")
                        col_idx += cell.grid_span
                        continue
                    else:
                        cell_set.add(cell._tc)

                    spanned_idx = row_idx
                    spanned_tc: Optional[CT_Tc] = cell._tc
                    while spanned_tc == cell._tc:
                        spanned_idx += 1
                        spanned_tc = (
                            table.rows[spanned_idx].cells[col_idx]._tc
                            if spanned_idx < num_rows
                            else None
                        )
                    _log.debug(f"  spanned before row {spanned_idx}")

                    table_cell = TableCell(
                        text=cell.text,
                        row_span=spanned_idx - row_idx,
                        col_span=cell.grid_span,
                        start_row_offset_idx=row.grid_cols_before + row_idx,
                        end_row_offset_idx=row.grid_cols_before + spanned_idx,
                        start_col_offset_idx=col_idx,
                        end_col_offset_idx=col_idx + cell.grid_span,
                        col_header=False,
                        row_header=False,
                    )
                    data.table_cells.append(table_cell)
                    col_idx += cell.grid_span

        level = self.get_level()
        doc.add_table(data=data, parent=self.parents[level - 1])
        return

    def handle_pictures(
        self, 
        docx_obj: DocxDocument, 
        drawing_blip: Any, 
        doc: DoclingDocument,
        parent: Optional[NodeItem] = None
    ) -> None:
        def get_docx_image(blip):
            rId = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if rId and rId in docx_obj.part.rels:
                # Access the image part using the relationship ID
                image_part = docx_obj.part.rels[rId].target_part
                return image_part.blob  # Get the binary image data
            return None

        # Use provided parent or get current level's parent
        if parent is None:
            level = self.get_level()
            parent = self.parents[level - 1]
        
        # Process all blips in the drawing_blip
        for blip in drawing_blip:
            # Open the BytesIO object with PIL to create an Image
            try:
                image_data = get_docx_image(blip)
                if image_data:
                    image_bytes = BytesIO(image_data)
                    pil_image = Image.open(image_bytes)
                    doc.add_picture(
                        parent=parent,
                        image=ImageRef.from_pil(image=pil_image, dpi=72),
                        caption=None,
                    )
            except (UnidentifiedImageError, OSError) as e:
                _log.warning(f"Warning: image cannot be loaded by Pillow: {str(e)}")
                doc.add_picture(
                    parent=parent,
                    caption=None,
                )
        return

    def extract_all_images(self, docx_obj: DocxDocument) -> list[str]:
        """Extract all image relationships IDs from the document.
        
        This method can be used to validate that all images are being properly processed.
        
        Returns:
            A list of image relationship IDs found in the document.
        """
        # Define the namespaces
        namespaces = self.xml_namespaces
        
        # Get the root element of the document
        root = docx_obj.element
        
        # Find all image references
        image_refs = []
        
        # Look for embedded images
        for drawing in root.findall('.//w:drawing', namespaces):
            for blip in drawing.findall('.//a:blip', namespaces):
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id:
                    image_refs.append(embed_id)
        
        # Also check for images in tables
        for table in root.findall('.//w:tbl', namespaces):
            for drawing in table.findall('.//w:drawing', namespaces):
                for blip in drawing.findall('.//a:blip', namespaces):
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        image_refs.append(embed_id)
        
        return image_refs
    
    def handle_textbox(
        self,
        txbx_content: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        """Process content inside a text box."""
        level = self.get_level()
        
        # Create a container for the text box content
        textbox_node = doc.add_group(
            label=GroupLabel.SECTION, 
            name="textbox", 
            parent=self.parents[level - 1]
        )
        
        # Save the current parent state
        original_parents = self.parents.copy()
        
        # Set the textbox node as the current parent for its contents
        self.parents[level] = textbox_node
        
        # Process paragraphs inside the text box
        for element in txbx_content:
            tag_name = etree.QName(element).localname
            if tag_name == "p":
                self.handle_text_elements(element, docx_obj, doc)
            elif tag_name == "tbl":
                try:
                    self.handle_tables(element, docx_obj, doc)
                except Exception as e:
                    _log.debug(f"Could not parse a table in textbox, broken docx table: {str(e)}")
            elif tag_name == "sdt":
                sdt_content = element.find(".//w:sdtContent", namespaces=self.xml_namespaces)
                if sdt_content is not None:
                    paragraphs = sdt_content.findall(".//w:p", namespaces=self.xml_namespaces)
                    for p in paragraphs:
                        self.handle_text_elements(p, docx_obj, doc)
        
        # Restore original parent state after processing textbox contents
        self.parents = original_parents

    def export_table_of_contents(
        self,
        include_title: bool = True,
        delim: str = "\n",
        escaping_underscores: bool = True,
        indent: int = 2,
        text_width: int = -1,
        page_no: Optional[int] = None,
    ) -> str:
        """Export just the table of contents in Markdown format.
        
        Args:
            include_title: Whether to include the "Table of Contents" title
            delim: Delimiter to use when concatenating the TOC parts
            escaping_underscores: Whether to escape underscores in the text content
            indent: The indent in spaces per level for the nested list items
            text_width: The width to wrap text at (-1 for no wrapping)
            page_no: Optional page number to filter content by
            
        Returns:
            str: The TOC in Markdown format
        """
        # Start with title if requested
        toc = ["# Table of Contents"] if include_title else []
        
        # Track previously seen headings to avoid duplicates
        seen_headings = set()
        
        # Process all heading items
        for item, level in self.iterate_items(self.body, with_groups=True, page_no=page_no):
            
            # Only process heading items
            if (isinstance(item, TextItem) and item.label == DocItemLabel.SECTION_HEADER) or \
            isinstance(item, SectionHeaderItem):
                heading_text = item.text.strip()
                if not heading_text:
                    continue
                    
                # Skip duplicates
                if heading_text in seen_headings:
                    continue
                seen_headings.add(heading_text)
                
                # Get heading level - handle special cases
                heading_level = getattr(item, 'level', 1)
                if heading_level < 1:  # Ensure valid level
                    heading_level = 1
                    
                # Create indentation based on heading level
                level_indent = " " * ((heading_level - 1) * indent)
                
                # Create anchor link from heading text
                link_text = heading_text.lower()
                link_text = re.sub(r'[^\w\s-]', '', link_text)
                link_text = re.sub(r'[\s_]+', '-', link_text)
                
                # Escape underscores in display text if needed
                display_text = heading_text
                if escaping_underscores:
                    display_text = display_text.replace('_', '\\_')
                
                # Add entry to TOC
                toc.append(f"{level_indent}- [{display_text}]") #(#{link_text})
        
        # Add a blank line after the TOC if there are entries
        if len(toc) > (1 if include_title else 0):
            toc.append("")
        
        # Join with delimiter
        toc_content = delim.join(toc)
        
        # Apply text width if specified
        if text_width > 0:
            wrapped_lines = []
            for line in toc_content.split('\n'):
                # Don't wrap headings or list items, only their content
                if re.match(r'^(\s*-|\s*#)', line):
                    # Find the position after the list marker or heading symbol
                    match = re.match(r'^(\s*-\s+|\s*#+\s+)', line)
                    if match:
                        prefix_len = len(match.group(0))
                        prefix = line[:prefix_len]
                        content = line[prefix_len:]
                        # Wrap the content part with proper indentation
                        wrapped_content = textwrap.fill(
                            content, 
                            width=text_width,
                            initial_indent='',
                            subsequent_indent=' ' * prefix_len
                        )
                        wrapped_lines.append(prefix + wrapped_content)
                    else:
                        wrapped_lines.append(line)
                else:
                    wrapped_lines.append(textwrap.fill(line, width=text_width))
            toc_content = '\n'.join(wrapped_lines)
        
        return toc_content

    def extract_document_metadata(self) -> tuple[Optional[datetime], Optional[datetime], Optional[str]]:
        """
        Extract creation date, last modified date, and author from the DOCX metadata.
        
        Returns:
            Tuple of (creation_date, modified_date, author) where dates are datetime objects,
            author is a string, or None if not available
        """
        created_date = None
        modified_date = None
        author = None
        
        try:
            if self.docx_obj is not None:
                # Access the core properties part
                core_props = self.docx_obj.core_properties
                
                # Extract creation date if available
                if hasattr(core_props, 'created') and core_props.created:
                    created_date = core_props.created
                
                # Extract modification date if available
                if hasattr(core_props, 'modified') and core_props.modified:
                    modified_date = core_props.modified
                    
                # Extract author if available
                if hasattr(core_props, 'author') and core_props.author:
                    author = core_props.author
                    
        except Exception as e:
            _log.warning(f"Error extracting document metadata: {str(e)}")
            
        return created_date, modified_date, author
    
    def handle_diagram(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
        parent: Optional[NodeItem] = None
    ) -> None:
        """Process SmartArt diagrams in the document with more detailed extraction."""
        # Use provided parent or get current level's parent
        if parent is None:
            level = self.get_level()
            parent = self.parents[level - 1]
        
        # Extract diagram relationships
        diagram_rel = element.find(".//dgm:relIds", namespaces=self.xml_namespaces)
        if diagram_rel is not None:
            # Get all relationship IDs
            dm_rel = diagram_rel.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}dm")  # Data model
            lo_rel = diagram_rel.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lo")  # Layout
            qs_rel = diagram_rel.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}qs")  # Quick style
            cs_rel = diagram_rel.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}cs")  # Color style
            
            diagram_data = None
            diagram_shapes = []
            
            # Try to access the actual diagram data through relationships
            try:
                # First try the data model
                if dm_rel and dm_rel in docx_obj.part.rels:
                    diagram_part = docx_obj.part.rels[dm_rel].target_part
                    if hasattr(diagram_part, 'blob'):
                        # Extract text from data model
                        diagram_shapes = self.extract_diagram_shapes(diagram_part.blob)
                
                # If no text found in data model, try the drawing XML
                if not diagram_shapes and element is not None:
                    # Extract any embedded drawing content (like in the paste-2.txt example)
                    xpath_drawing = XPath(".//dsp:drawing", namespaces=self.xml_namespaces)
                    drawing_elements = xpath_drawing(element)
                    
                    if drawing_elements:
                        for drawing in drawing_elements:
                            # Extract shapes from the drawing
                            xpath_shapes = XPath(".//dsp:sp", namespaces=self.xml_namespaces)
                            shapes = xpath_shapes(drawing)
                            
                            for shape in shapes:
                                shape_texts = []
                                # Find text body
                                text_body = shape.find(".//dsp:txBody", namespaces=self.xml_namespaces)
                                if text_body is not None:
                                    # Extract paragraphs
                                    paragraphs = text_body.findall(".//a:p", namespaces=self.xml_namespaces)
                                    for p in paragraphs:
                                        para_texts = []
                                        # Extract text runs
                                        text_runs = p.findall(".//a:r", namespaces=self.xml_namespaces)
                                        for r in text_runs:
                                            t_elem = r.find(".//a:t", namespaces=self.xml_namespaces)
                                            if t_elem is not None and t_elem.text:
                                                para_texts.append(t_elem.text)
                                        
                                        # Find line breaks
                                        line_breaks = p.findall(".//a:br", namespaces=self.xml_namespaces)
                                        if line_breaks:
                                            # If there are line breaks, handle them
                                            shape_texts.append(" ".join(para_texts))
                                            para_texts = []  # Start a new paragraph after line break
                                        elif para_texts:
                                            shape_texts.append(" ".join(para_texts))
                                
                                if shape_texts:
                                    diagram_shapes.append(shape_texts)
                
                # Process diagram content if found
                if diagram_shapes:
                    # Create a section for the diagram
                    diagram_section = doc.add_group(
                        label=GroupLabel.SECTION,
                        name="diagram",
                        parent=parent
                    )
                    """
                    # Add a header for the diagram
                    doc.add_heading(
                        parent=diagram_section,
                        text="SmartArt Diagram",
                        level=3
                    )
                    """
                    doc.add_text(
                        label=DocItemLabel.PARAGRAPH,
                        parent=diagram_section,
                        text="Diagram:"
                    )
                    
                    # Add shape content as paragraphs
                    for shape_idx, shape_texts in enumerate(diagram_shapes):
                        # Merge text fragments within each shape
                        merged_texts = self.merge_diagram_text(shape_texts)
                        
                        for text in merged_texts:
                            if text.strip():
                                doc.add_text(
                                    label=DocItemLabel.PARAGRAPH,
                                    parent=diagram_section,
                                    text=text
                                )
                    
                    return
                    
            except Exception as e:
                _log.warning(f"Error extracting diagram content: {str(e)}")
            
            # Fallback: Create a generic placeholder
            doc.add_text(
                label=DocItemLabel.PARAGRAPH,
                parent=parent,
                text="[SmartArt Diagram]"
            )
        
        return

    def extract_diagram_shapes(self, xml_content: bytes) -> List[List[str]]:
        """Extract text grouped by shapes from diagram XML."""
        try:
            from lxml import etree
            root = etree.fromstring(xml_content)
            
            shapes = []
            # Find all diagram nodes/shapes
            shape_elements = root.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/diagram}pt")
            
            for shape in shape_elements:
                shape_texts = []
                # Find text elements in this shape
                text_elements = shape.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}t")
                for text_elem in text_elements:
                    if text_elem.text and text_elem.text.strip():
                        shape_texts.append(text_elem.text.strip())
                
                if shape_texts:
                    shapes.append(shape_texts)
            
            return shapes
        except Exception as e:
            _log.warning(f"Error parsing diagram XML: {str(e)}")
            return []

    def merge_diagram_text(self, text_fragments: List[str]) -> List[str]:
        """
        Merge text fragments into logical segments, starting a new segment when:
        1. A fragment starts with a capital letter (and isn't continuing a sentence)
        2. There's a line break or other clear separator
        
        Args:
            text_fragments: List of raw text fragments extracted from diagram XML
            
        Returns:
            List of merged text segments
        """
        if not text_fragments:
            return []
        
        merged_segments = []
        current_segment = ""
        
        for fragment in text_fragments:
            # Skip empty fragments
            if not fragment:
                continue
                
            # Check if this fragment should start a new segment
            starts_with_capital = fragment[0].isupper() if fragment else False
            is_continuation = False
            
            # Check if it's just continuing a sentence (e.g. after comma, etc.)
            if starts_with_capital and current_segment and current_segment.strip():
                last_char = current_segment.strip()[-1]
                # If last segment ended with punctuation that doesn't end a sentence
                if last_char in [',', ';', ':', '(', '-', '–', '—']:
                    is_continuation = True
                    
            # Start a new segment if:
            # - Current fragment starts with capital (not a continuation)
            # - Current segment is not empty
            if starts_with_capital and current_segment and not is_continuation:
                merged_segments.append(current_segment.strip())
                current_segment = fragment
            # Otherwise, append to the current segment with a space
            else:
                if current_segment:
                    # Add appropriate spacing between fragments
                    if current_segment.endswith(('-', '–', '—')):
                        # No space needed after hyphens
                        current_segment += fragment
                    else:
                        current_segment += " " + fragment
                else:
                    current_segment = fragment
        
        # Add the last segment if not empty
        if current_segment:
            merged_segments.append(current_segment.strip())
        
        return merged_segments